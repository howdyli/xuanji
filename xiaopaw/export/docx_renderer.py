"""Render Markdown text to DOCX byte stream."""

from __future__ import annotations

import io
import markdown
from xml.etree import ElementTree as ET

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def render_markdown_to_docx(markdown_text: str) -> bytes:
    """将 Markdown 文本渲染为 DOCX 字节流。"""
    html_body = markdown.markdown(
        markdown_text, extensions=["tables", "fenced_code", "nl2br"]
    )
    root = ET.fromstring(f"<root>{html_body}</root>")

    doc = Document()
    for elem in root:
        _process_element(doc, elem)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── element dispatch ─────────────────────────────────────────────

def _process_element(doc: Document, elem: ET.Element) -> None:
    tag = elem.tag
    if tag in ("h1", "h2", "h3"):
        doc.add_heading(elem.text or "", level=int(tag[1]))
    elif tag == "p":
        _add_formatted_paragraph(doc, elem)
    elif tag == "pre":
        code = elem.find("code")
        _add_code_block(doc, code.text if code is not None and code.text else "")
    elif tag == "blockquote":
        _add_blockquote(doc, elem)
    elif tag == "hr":
        doc.add_paragraph("─" * 50)
    elif tag == "table":
        _add_table(doc, elem)
    elif tag == "ul":
        _add_list(doc, elem, ordered=False)
    elif tag == "ol":
        _add_list(doc, elem, ordered=True)


# ── code block ───────────────────────────────────────────────────

def _add_code_block(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "F5F5F5")
    shading.set(qn("w:val"), "clear")
    p.paragraph_format.element.get_or_add_pPr().append(shading)


# ── blockquote ───────────────────────────────────────────────────

def _add_blockquote(doc: Document, elem: ET.Element) -> None:
    text = _get_text(elem)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    run = p.add_run(text)
    run.italic = True
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


# ── table ────────────────────────────────────────────────────────

def _add_table(doc: Document, table_elem: ET.Element) -> None:
    rows = table_elem.findall(".//tr")
    if not rows:
        return
    cols = len(rows[0].findall("./th") or rows[0].findall("./td"))
    if cols == 0:
        return
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        cells = row.findall("./th") or row.findall("./td")
        for j, cell in enumerate(cells):
            if j < cols:
                table.rows[i].cells[j].text = cell.text or ""


# ── lists ────────────────────────────────────────────────────────

def _add_list(doc: Document, elem: ET.Element, *, ordered: bool) -> None:
    style = "List Number" if ordered else "List Bullet"
    for i, li in enumerate(elem.findall("li")):
        text = _get_text(li)
        prefix = f"{i + 1}. " if ordered else ""
        doc.add_paragraph(f"{prefix}{text}", style=style)


# ── inline formatting ───────────────────────────────────────────

def _add_formatted_paragraph(doc: Document, elem: ET.Element) -> None:
    p = doc.add_paragraph()
    _process_inline(p, elem)


def _process_inline(paragraph, elem: ET.Element) -> None:
    if elem.text:
        paragraph.add_run(elem.text)
    for child in elem:
        if child.tag in ("strong", "b"):
            run = paragraph.add_run(child.text or "")
            run.bold = True
        elif child.tag in ("em", "i"):
            run = paragraph.add_run(child.text or "")
            run.italic = True
        elif child.tag == "code":
            run = paragraph.add_run(child.text or "")
            run.font.name = "Consolas"
            run.font.size = Pt(9)
        elif child.tag == "a":
            run = paragraph.add_run(child.text or "")
            run.font.color.rgb = RGBColor(0x7C, 0x6A, 0xF4)
        else:
            paragraph.add_run(child.text or "")
        if child.tail:
            paragraph.add_run(child.tail)


# ── helpers ──────────────────────────────────────────────────────

def _get_text(elem: ET.Element) -> str:
    """Recursively collect all text from an element."""
    return "".join(elem.itertext())
